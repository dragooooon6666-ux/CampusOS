"""写作中心 API"""

import json

from fastapi import APIRouter, HTTPException

from backend.database import get_db
from backend.services.template_service import (
    get_form_fields,
    list_templates,
    get_template,
    create_template,
)
from backend.services.writing_engine import generate_document, regenerate_section

router = APIRouter(prefix="/api/writing", tags=["writing"])


# ── 表单字段 ──

@router.get("/form-fields/{doc_type}")
def form_fields(doc_type: str):
    return {"doc_type": doc_type, "fields": get_form_fields(doc_type)}


# ── 生成文档 ──

@router.post("/generate")
def generate(data: dict):
    # 仅更新项目关联
    if data.get("doc_type") == "_link_":
        doc_id = data.get("_doc_id")
        project_id = data.get("project_id")
        if not doc_id or not project_id:
            raise HTTPException(400, "缺少参数")
        db = get_db()
        db.execute("UPDATE documents SET project_id = ? WHERE id = ?", (project_id, doc_id))
        db.commit()
        row = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
        db.close()
        return dict(row) if row else {}

    doc_type = data.get("doc_type", "")
    if not doc_type:
        raise HTTPException(400, "缺少 doc_type")
    form_data = data.get("form_data", {})
    template_id = data.get("template_id")

    try:
        result = generate_document(doc_type, form_data, template_id)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"AI 生成失败: {e}")

    # 保存到数据库
    db = get_db()
    project_id = data.get("project_id")
    cur = db.execute(
        "INSERT INTO documents (project_id, doc_type, title, content, ai_provider, template_id) "
        "VALUES (?, ?, ?, ?, ?, ?)",
        (project_id, doc_type, result["title"], result["content"],
         result["provider"], template_id),
    )
    db.commit()
    doc_id = cur.lastrowid
    row = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
    db.close()

    return dict(row)


# ── 重新生成段落 ──

@router.post("/regenerate")
def regenerate(data: dict):
    doc_id = data.get("doc_id")
    section_key = data.get("section_key", "")
    section_label = data.get("section_label", "")
    feedback = data.get("feedback", "")

    if not doc_id:
        raise HTTPException(400, "缺少 doc_id")

    db = get_db()
    row = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
    if not row:
        db.close()
        raise HTTPException(404, "文档不存在")

    # 直接保存（手动编辑后保存）
    if section_key == '_save' and feedback.startswith('__FULL_REPLACE__:'):
        new_content = feedback.split('__FULL_REPLACE__:', 1)[1]
    else:
        try:
            new_content = regenerate_section(
                row["doc_type"], row["content"],
                section_key, section_label, feedback,
            )
        except Exception as e:
            db.close()
            raise HTTPException(500, f"重新生成失败: {e}")

    db.execute("UPDATE documents SET content = ? WHERE id = ?", (new_content, doc_id))
    db.commit()
    row = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
    db.close()

    return dict(row)


# ── 导出 docx ──

@router.get("/documents/{doc_id}/export")
def export_docx(doc_id: int):
    """将文档导出为 docx 文件下载"""
    from fastapi.responses import Response
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import io, re

    def _set_run(run, font_name, size, bold=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = size
        run.bold = bold

    db = get_db()
    row = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
    db.close()
    if not row:
        raise HTTPException(404, "文档不存在")

    doc = Document()
    # 全局默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    content = row["content"]
    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run(p.add_run(line[3:]), '宋体', Pt(16), bold=True)
        elif line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run(p.add_run(line[2:]), '宋体', Pt(18), bold=True)
        else:
            p = doc.add_paragraph()
            _set_run(p.add_run(line), '宋体', Pt(12))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = re.sub(r'[<>:"/\\|?*]', '', row["title"] or "文档")
    filename = f"{safe_title}.docx"

    # 保存到 output 文件夹
    from pathlib import Path
    from backend.services.archiver import OUTPUT_DIR, DOC_TYPE_GROUP
    group = DOC_TYPE_GROUP.get(row["doc_type"], "其他分类")
    dest_dir = OUTPUT_DIR / "写作导出" / group
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest_path = dest_dir / filename
    counter = 1
    while dest_path.exists():
        dest_path = dest_dir / f"{safe_title}({counter}).docx"
        counter += 1
    import shutil
    buf.seek(0)
    with open(dest_path, "wb") as f:
        shutil.copyfileobj(io.BytesIO(buf.getvalue()), f)

    # 记录到 files 表并关联项目
    import hashlib
    content_hash = hashlib.md5(buf.getvalue()).hexdigest()
    db2 = get_db()
    cur = db2.execute(
        """INSERT INTO files (original_name, stored_name, original_path, stored_path,
           extension, size_bytes, doc_type, title, doc_date, content_hash, ai_analyzed)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0)""",
        (filename, dest_path.name, str(dest_path), str(dest_path),
         ".docx", len(buf.getvalue()), row["doc_type"], safe_title, "", content_hash),
    )
    file_id = cur.lastrowid
    db2.commit()
    db2.close()

    buf.seek(0)
    from urllib.parse import quote
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(safe_title)}.docx"},
    )


# ── 文档列表 ──

@router.get("/documents")
def list_documents(project_id: int | None = None):
    db = get_db()
    if project_id:
        rows = db.execute(
            "SELECT * FROM documents WHERE project_id = ? ORDER BY created_at DESC",
            (project_id,),
        ).fetchall()
    else:
        rows = db.execute("SELECT * FROM documents ORDER BY created_at DESC").fetchall()
    db.close()
    return [dict(r) for r in rows]


# ── 模板 API ──

@router.get("/templates")
def get_templates(doc_type: str | None = None):
    return list_templates(doc_type)


@router.post("/templates")
def add_template(data: dict):
    name = data.get("name", "")
    doc_type = data.get("doc_type", "")
    style = data.get("style", "通用")
    sections = data.get("sections", [])
    if not name or not doc_type or not sections:
        raise HTTPException(400, "缺少 name/doc_type/sections")
    return create_template(name, doc_type, style, sections, data.get("formatting"))


@router.post("/templates/import")
def import_template(data: dict):
    """从文档文件导入模板结构（data 中传 template_name 和 sections）"""
    name = data.get("name", "导入模板")
    doc_type = data.get("doc_type", "其他")
    sections = data.get("sections", [])
    if not sections:
        raise HTTPException(400, "sections 不能为空")
    return create_template(name, doc_type, "自定义", sections)


@router.post("/templates/extract")
def extract_template(data: dict):
    """AI 提取文档结构 → 返回 sections（不保存）"""
    text = data.get("text", "")
    if not text:
        raise HTTPException(400, "缺少文档文本")

    from backend.services.writing_engine import _call_ai
    prompt = f"""分析以下文档的结构，提取其章节划分，生成模板 sections 数组。

每个 section: {{"key": "英文标识", "label": "中文标签", "ai_role": "AI生成该段落的角色描述"}}

只返回 JSON 数组，不要其他内容。文档内容：

{text[:2000]}"""

    try:
        raw = _call_ai(
            "你是文档结构分析助手。提取章节结构，只输出 JSON 数组。",
            prompt,
            temperature=0.3,
        )
        import re
        m = re.search(r'\[.*\]', raw, re.DOTALL)
        sections = json.loads(m.group()) if m else []
        return {"sections": sections}
    except Exception as e:
        raise HTTPException(500, f"AI 提取失败: {e}")
